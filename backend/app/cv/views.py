import io
import os

from django.http import FileResponse
from rest_framework import status
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.parsers import MultiPartParser, JSONParser
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from .models import CV, TailoredCV, CoverLetter
from .serializers import CVSerializer, TailoredCVSerializer, CoverLetterSerializer
from .cv_parser import extract_text_from_pdf, parse_cv_text
from .cv_tailor_agent import CVTailorAgent
from .cover_letter_agent import CoverLetterAgent

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB
VALID_CHANGE_STATUSES = {'pending', 'accepted', 'rejected'}


def _update_change_statuses(change_set: dict, changes: list) -> dict:
    if not isinstance(change_set, dict):
        change_set = {}

    indexed_changes = {}
    for section_name in ('skills', 'experience', 'education'):
        section_changes = change_set.get(section_name, [])
        if isinstance(section_changes, list):
            indexed_changes[section_name] = section_changes

    for change in changes:
        section_name = change.get('section')
        change_id = change.get('id')
        new_status = change.get('status')

        if section_name not in indexed_changes:
            raise ValueError(f'Unknown section "{section_name}".')
        if new_status not in VALID_CHANGE_STATUSES:
            raise ValueError(f'Invalid status "{new_status}".')

        matched = False
        for item in indexed_changes[section_name]:
            if isinstance(item, dict) and item.get('id') == change_id:
                item['status'] = new_status
                matched = True
                break

        if not matched:
            raise ValueError(f'Change "{change_id}" not found in section "{section_name}".')

    return change_set


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@parser_classes([MultiPartParser])
def upload_cv(request):
    """Upload a PDF CV, extract text, and return parsed data."""
    file = request.FILES.get('file')

    if not file:
        return Response({'error': 'No file provided.'}, status=status.HTTP_400_BAD_REQUEST)

    # Validate extension
    ext = os.path.splitext(file.name)[1].lower()
    if ext != '.pdf':
        return Response({'error': 'Only .pdf files are accepted.'}, status=status.HTTP_400_BAD_REQUEST)

    # Validate size
    if file.size > MAX_FILE_SIZE:
        return Response(
            {'error': f'File exceeds the 5 MB limit (got {file.size} bytes).'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    # Extract & parse
    raw_text = extract_text_from_pdf(file)
    parsed = parse_cv_text(raw_text)

    # Persist
    cv = CV.objects.create(
        user=request.user,
        file=file,
        extracted_name=parsed['name'],
        extracted_skills=parsed['skills'],
        extracted_experience=parsed['experience'],
        extracted_education=parsed['education'],
    )

    return Response(CVSerializer(cv).data, status=status.HTTP_201_CREATED)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def cv_me(request):
    """Return the current user's most recent CV, or 404 if none."""
    cv = request.user.cvs.first()
    if not cv:
        return Response({'error': 'No CV found.'}, status=status.HTTP_404_NOT_FOUND)
    return Response(CVSerializer(cv).data)


@api_view(['GET', 'PUT', 'PATCH'])
@permission_classes([IsAuthenticated])
@parser_classes([JSONParser])
def cv_detail(request, pk):
    """Retrieve or update extracted fields of a CV owned by the current user."""
    try:
        cv = CV.objects.get(pk=pk, user=request.user)
    except CV.DoesNotExist:
        return Response({'error': 'CV not found.'}, status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        return Response(CVSerializer(cv).data)

    # PUT / PATCH
    partial = request.method == 'PATCH'
    serializer = CVSerializer(cv, data=request.data, partial=partial)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# ───────────────────────────────────────────────
#  CV Tailoring (Agent 2)
# ───────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def tailor_cv(request):
    """Tailor the user's CV for a specific job listing.

    Expects JSON body: { job_title, job_company?, job_description }
    """
    job_title = request.data.get('job_title', '').strip()
    job_description = request.data.get('job_description', '').strip()
    job_company = request.data.get('job_company', '').strip()

    if not job_title or not job_description:
        return Response(
            {'error': 'job_title and job_description are required.'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    cv = request.user.cvs.first()
    if not cv:
        return Response(
            {'error': 'Please upload a CV first.'},
            status=status.HTTP_404_NOT_FOUND,
        )

    cv_data = {
        'skills': cv.extracted_skills or [],
        'experience': cv.extracted_experience or [],
        'education': cv.extracted_education or [],
    }

    agent = CVTailorAgent()
    tailored = agent.tailor(cv_data, {
        'title': job_title,
        'company': job_company,
        'description': job_description,
    })

    tailored_cv = TailoredCV.objects.create(
        user=request.user,
        original_cv=cv,
        job_title=job_title,
        job_company=job_company,
        job_description=job_description,
        original_skills=cv_data['skills'],
        original_experience=cv_data['experience'],
        original_education=cv_data['education'],
        change_set=tailored['change_set'],
        tailored_skills=tailored['tailored_skills'],
        tailored_experience=tailored['tailored_experience'],
        tailored_education=tailored['tailored_education'],
    )

    return Response(TailoredCVSerializer(tailored_cv).data, status=status.HTTP_201_CREATED)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def tailored_cv_list(request):
    """Return all tailored CVs for the authenticated user."""
    tailored = TailoredCV.objects.filter(user=request.user)
    return Response(TailoredCVSerializer(tailored, many=True).data)


@api_view(['GET', 'PATCH'])
@permission_classes([IsAuthenticated])
@parser_classes([JSONParser])
def tailored_cv_detail(request, pk):
    """Retrieve or update review statuses for a tailored CV."""
    try:
        tailored_cv = TailoredCV.objects.get(pk=pk, user=request.user)
    except TailoredCV.DoesNotExist:
        return Response({'error': 'Tailored CV not found.'}, status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        return Response(TailoredCVSerializer(tailored_cv).data)

    changes = request.data.get('changes')
    if not isinstance(changes, list) or not changes:
        return Response(
            {'error': 'changes must be a non-empty list.'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    try:
        tailored_cv.change_set = _update_change_statuses(tailored_cv.change_set, changes)
    except ValueError as exc:
        return Response({'error': str(exc)}, status=status.HTTP_400_BAD_REQUEST)

    tailored_cv.save(update_fields=['change_set'])
    return Response(TailoredCVSerializer(tailored_cv).data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_tailored_cv(request, pk):
    """Download a tailored CV as a PDF file."""
    try:
        tailored_cv = TailoredCV.objects.get(pk=pk, user=request.user)
    except TailoredCV.DoesNotExist:
        return Response({'error': 'Tailored CV not found.'}, status=status.HTTP_404_NOT_FOUND)

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        return Response(
            {'error': 'PDF generation library not available.'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CVTitle',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        spaceAfter=20,
        fontSize=18
    )

    section_heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        spaceBefore=15,
        spaceAfter=5,
        fontSize=14,
        textColor='#333333'
    )

    body_style = styles['Normal']
    body_style.fontSize = 11
    body_style.spaceAfter = 6

    bullet_style = ParagraphStyle(
        'BulletItem',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=4,
        leftIndent=15
    )

    story = []

    # Title / Name
    name = tailored_cv.original_cv.extracted_name if tailored_cv.original_cv.extracted_name else "Tailored CV"
    story.append(Paragraph(name, title_style))

    job_target = f"Target Role: {tailored_cv.job_title}"
    if tailored_cv.job_company:
        job_target += f" at {tailored_cv.job_company}"

    target_style = ParagraphStyle(
        'JobTarget',
        parent=body_style,
        alignment=TA_CENTER,
        textColor='#666666'
    )
    story.append(Paragraph(job_target, target_style))
    story.append(Spacer(1, 15))

    def get_final_items(section_key, tailored_items):
        changes = tailored_cv.change_set.get(section_key, [])
        final = []
        for i, item in enumerate(tailored_items):
            change_id = f"{section_key}-{i}"
            change = next((c for c in changes if isinstance(c, dict) and c.get('id') == change_id), None)

            if change and change.get('status') == 'rejected':
                final.append(change.get('before', item))
            else:
                final.append(item)
        return final

    # Skills
    skills = get_final_items('skills', tailored_cv.tailored_skills)
    if skills:
        story.append(Paragraph('Skills', section_heading_style))
        skill_items = [ListItem(Paragraph(skill, bullet_style)) for skill in skills]
        story.append(ListFlowable(skill_items, bulletType='bullet', start='circle'))
        story.append(Spacer(1, 10))

    # Experience
    experience = get_final_items('experience', tailored_cv.tailored_experience)
    if experience:
        story.append(Paragraph('Experience', section_heading_style))
        exp_items = [ListItem(Paragraph(exp, bullet_style)) for exp in experience]
        story.append(ListFlowable(exp_items, bulletType='bullet', start='circle'))
        story.append(Spacer(1, 10))

    # Education
    education = get_final_items('education', tailored_cv.tailored_education)
    if education:
        story.append(Paragraph('Education', section_heading_style))
        edu_items = [ListItem(Paragraph(edu, bullet_style)) for edu in education]
        story.append(ListFlowable(edu_items, bulletType='bullet', start='circle'))
        story.append(Spacer(1, 10))

    doc.build(story)
    buf.seek(0)

    safe_title = ''.join(c for c in tailored_cv.job_title if c.isalnum() or c in ' -_')[:50]
    filename = f'Tailored_CV_{safe_title}.pdf'

    return FileResponse(
        buf,
        as_attachment=True,
        filename=filename,
        content_type='application/pdf',
    )


# ───────────────────────────────────────────────
#  Cover Letter Generation (Agent 2)
# ───────────────────────────────────────────────


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_cover_letter(request):
    """Generate a cover letter for a specific job listing.

    Expects JSON body: { job_title, job_company?, job_description }
    """
    job_title = request.data.get('job_title', '').strip()
    job_description = request.data.get('job_description', '').strip()
    job_company = request.data.get('job_company', '').strip()

    if not job_title or not job_description:
        return Response(
            {'error': 'job_title and job_description are required.'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    cv = request.user.cvs.first()
    if not cv:
        return Response(
            {'error': 'Please upload a CV first.'},
            status=status.HTTP_404_NOT_FOUND,
        )

    cv_data = {
        'name': cv.extracted_name or '',
        'skills': cv.extracted_skills or [],
        'experience': cv.extracted_experience or [],
        'education': cv.extracted_education or [],
    }

    agent = CoverLetterAgent()
    body = agent.generate(cv_data, {
        'title': job_title,
        'company': job_company,
        'description': job_description,
    })

    cover_letter = CoverLetter.objects.create(
        user=request.user,
        original_cv=cv,
        job_title=job_title,
        job_company=job_company,
        job_description=job_description,
        body=body,
    )

    return Response(CoverLetterSerializer(cover_letter).data, status=status.HTTP_201_CREATED)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def cover_letter_list(request):
    """Return all cover letters for the authenticated user."""
    letters = CoverLetter.objects.filter(user=request.user)
    return Response(CoverLetterSerializer(letters, many=True).data)


@api_view(['GET', 'PATCH'])
@permission_classes([IsAuthenticated])
def cover_letter_detail(request, pk):
    """Retrieve or update a cover letter owned by the current user."""
    try:
        letter = CoverLetter.objects.get(pk=pk, user=request.user)
    except CoverLetter.DoesNotExist:
        return Response({'error': 'Cover letter not found.'}, status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        return Response(CoverLetterSerializer(letter).data)

    # PATCH — only body is editable
    serializer = CoverLetterSerializer(letter, data=request.data, partial=True)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_cover_letter(request, pk):
    """Download a cover letter as a .docx file."""
    try:
        letter = CoverLetter.objects.get(pk=pk, user=request.user)
    except CoverLetter.DoesNotExist:
        return Response({'error': 'Cover letter not found.'}, status=status.HTTP_404_NOT_FOUND)

    from docx import Document
    from docx.shared import Pt

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add body paragraphs
    for paragraph_text in letter.body.split('\n'):
        if paragraph_text.strip():
            doc.add_paragraph(paragraph_text)
        else:
            doc.add_paragraph('')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = ''.join(c for c in letter.job_title if c.isalnum() or c in ' -_')[:50]
    filename = f'Cover_Letter_{safe_title}.docx'

    return FileResponse(
        buf,
        as_attachment=True,
        filename=filename,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
